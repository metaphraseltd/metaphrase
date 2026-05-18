#!/usr/bin/env python3
"""
Create Metaphrase Phase 2 Plan Word document
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
import copy

# Colours
NAVY = RGBColor(0x0F, 0x23, 0x40)
BLUE = RGBColor(0x1A, 0x5C, 0xB5)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
ROW_GREY_HEX = "F5F5F5"

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_num_pages(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'NUMPAGES'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_side, color_hex, size=12, space=0, val='single'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{border_side}')
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), str(space))
    border.set(qn('w:color'), color_hex)
    tcBorders.append(border)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph_border(para, color_hex, size=12, space=4, side='bottom'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), str(size))
    bdr.set(qn('w:space'), str(space))
    bdr.set(qn('w:color'), color_hex)
    pBdr.append(bdr)
    pPr.append(pBdr)

def set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_run_font(run, name):
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    run._r.get_or_add_rPr().insert(0, rFonts)

# ───────────────────────────────────────────────
# CREATE DOCUMENT
# ───────────────────────────────────────────────
doc = Document()

# Page setup: A4 with 2.5cm margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Content width in EMU
CONTENT_WIDTH_CM = 21 - 5  # 16cm
CONTENT_WIDTH_PT = Cm(16)

# ───────────────────────────────────────────────
# HEADER
# ───────────────────────────────────────────────
header = section.header
header.is_linked_to_previous = False

# Clear default header paragraph
for p in header.paragraphs:
    p._element.getparent().remove(p._element)

# Two-column table in header
# Total width = 16cm = ~9072 twips
# Left col: ~9000 twips, Right col: ~9000 twips
htable = header.add_table(1, 2, Cm(16))
htable.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove all borders
for row in htable.rows:
    for cell in row.cells:
        remove_cell_borders(cell)

# Left cell
left_cell = htable.cell(0, 0)
left_cell.width = Cm(8)

p1 = left_cell.paragraphs[0]
run_name = p1.add_run("Metaphrase")
run_name.bold = True
run_name.font.size = Pt(20)
run_name.font.color.rgb = NAVY

# Second line
p2 = left_cell.add_paragraph()
run_sub = p2.add_run("TRANSLATION & INTERPRETATION SERVICES")
run_sub.font.size = Pt(9)
run_sub.font.color.rgb = BLUE

# Third line
p3 = left_cell.add_paragraph()
run_co = p3.add_run("Company No. 16371031")
run_co.font.size = Pt(8)
run_co.font.color.rgb = GREY

# Right cell (left border blue)
right_cell = htable.cell(0, 1)
right_cell.width = Cm(8)
set_cell_border(right_cell, 'left', '1A5CB5', size=12)

contact_lines = [
    "metaphrase.uk",
    "info@metaphrase.uk",
    "+44 7852 408 467",
    "30 Victoria Road, Birmingham, B21 0SA"
]
first = True
for line in contact_lines:
    if first:
        rp = right_cell.paragraphs[0]
        first = False
    else:
        rp = right_cell.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = rp.add_run(line)
    r.font.size = Pt(8)
    r.font.color.rgb = GREY

# Blue paragraph border below the header table
# Add a blank paragraph with a blue bottom border
border_para = header.add_paragraph()
border_para.paragraph_format.space_before = Pt(4)
border_para.paragraph_format.space_after = Pt(0)
add_paragraph_border(border_para, '1A5CB5', size=12, space=1)

# ───────────────────────────────────────────────
# FOOTER
# ───────────────────────────────────────────────
footer = section.footer
footer.is_linked_to_previous = False

for p in footer.paragraphs:
    p._element.getparent().remove(p._element)

fp = footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

r1 = fp.add_run("Metaphrase Ltd | Phase 2: Translator Network Plan | Confidential | Page ")
r1.font.size = Pt(9)
r1.font.color.rgb = GREY

r_page = fp.add_run()
r_page.font.size = Pt(9)
r_page.font.color.rgb = GREY
add_page_number(r_page)

r2 = fp.add_run(" of ")
r2.font.size = Pt(9)
r2.font.color.rgb = GREY

r_total = fp.add_run()
r_total.font.size = Pt(9)
r_total.font.color.rgb = GREY
add_num_pages(r_total)

# ───────────────────────────────────────────────
# HELPER FUNCTIONS
# ───────────────────────────────────────────────

def add_heading(doc, text, level_pt=14, color=NAVY, gold_border=False, space_before=24, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(level_pt)
    r.font.color.rgb = color
    if gold_border:
        add_paragraph_border(p, 'C9A84C', size=8, space=4)
    return p

def add_heading2(doc, text, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p

def add_body(doc, text, space_before=4, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

def add_bullet(doc, text, space_before=2, space_after=2):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

# ───────────────────────────────────────────────
# TITLE PAGE
# ───────────────────────────────────────────────

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(32)
p_title.paragraph_format.space_after = Pt(8)
r = p_title.add_run("PHASE 2: BUILDING THE METAPHRASE TRANSLATOR NETWORK")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(8)
r2 = p_sub.add_run("Recruitment Strategy, Legal Framework & Growth Plan")
r2.font.size = Pt(13)
r2.font.color.rgb = BLUE

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_before = Pt(0)
p_date.paragraph_format.space_after = Pt(8)
r3 = p_date.add_run("May 2026 | Confidential")
r3.font.size = Pt(10)
r3.font.color.rgb = GREY
add_paragraph_border(p_date, 'C9A84C', size=8, space=4)

# Spacing paragraph after title
doc.add_paragraph()

# ───────────────────────────────────────────────
# SECTION 1 — EXECUTIVE SUMMARY
# ───────────────────────────────────────────────

add_heading(doc, "1. Executive Summary", gold_border=True)

add_body(doc, "Phase 2 marks the transition of Metaphrase Ltd from a solo linguist practice to a fully operational UK translation agency. The objective is to build a vetted network of freelance translators and interpreters across the UK's highest-demand language pairs, enabling Metaphrase to accept multi-language, high-volume, and concurrent projects — and to grow from a personal service into a scalable business capable of competing for corporate, NHS, and government contracts.")

add_body(doc, "This plan covers: the language pairs to target and why; the full legal framework for engaging UK freelancers; the recruitment form design and where to share it; how to integrate a Work With Us page on the website; a vetting and onboarding process; a suggested rate structure; and a 90-day action plan. Inshallah, with consistent execution this phase will have a network of 20+ vetted subcontractors active by the end of Q3 2026.")

# ───────────────────────────────────────────────
# SECTION 2 — LANGUAGE PAIRS
# ───────────────────────────────────────────────

add_heading(doc, "2. Language Pairs to Target", gold_border=True)

add_body(doc, "Selection is based on UK census demographics, demand from legal/medical/immigration sectors in the Midlands, and commercial value per word. Three tiers are defined: recruit immediately, recruit within 3 months, and grow into over 6–12 months. Languages with very small UK communities (Oriya, Kashmiri, Maithili, Sindhi etc.) are excluded at this stage — the cost of maintaining dormant subcontractor relationships outweighs the rare project opportunity.")

# Language pairs table
# Columns: Tier | Language Pair | Est. UK Speakers | Key Demand Sectors | Notes
col_widths = [Cm(1.8), Cm(3.2), Cm(2.5), Cm(4.5), Cm(4.0)]

rows_data = [
    # Header
    ["Tier", "Language Pair", "Est. UK Speakers", "Key Demand Sectors", "Notes"],
    # Tier 1
    ["Tier 1", "Urdu ↔ English", "400,000+", "Immigration (UKVI), courts, NHS, family law solicitors", "Kashif's own language — ensure 2–3 additional vetted backup translators"],
    ["Tier 1", "Punjabi ↔ English", "300,000+", "Crown & Magistrates Courts, NHS, police, immigration", "Very high court interpreter demand in West Midlands"],
    ["Tier 1", "Hindi ↔ English", "250,000+", "NHS, HR documents, business, immigration", "Overlaps with Urdu literacy in some documents"],
    ["Tier 1", "Arabic ↔ English", "200,000+", "Asylum cases, legal documents, business contracts", "High per-word rates; growing fast post-2015"],
    ["Tier 1", "Polish ↔ English", "800,000+", "Employment tribunals, housing, NHS, family courts", "Largest non-English community in UK; huge volume"],
    ["Tier 1", "Romanian ↔ English", "400,000+", "Employment, immigration, social services, NHS", "Very high Midlands demand specifically"],
    # Tier 2
    ["Tier 2", "Somali ↔ English", "100,000+", "Asylum, immigration, NHS, social services", "Severely underserved — premium rates achievable"],
    ["Tier 2", "Bengali / Sylheti ↔ English", "450,000+", "NHS, immigration, education, housing", "High NHS demand nationally"],
    ["Tier 2", "Gujarati ↔ English", "250,000+", "Business, property, inheritance, family law", "Strong Midlands community; often overlooked by agencies"],
    ["Tier 2", "Pashto / Dari ↔ English", "Growing rapidly", "Asylum, resettlement, immigration post-2021", "Currently severely undersupplied across the UK"],
    ["Tier 2", "Turkish ↔ English", "150,000+", "Legal, NHS, business", "Steady demand; not many agencies specialise"],
    ["Tier 2", "Mandarin ↔ English", "150,000+", "Business contracts, legal, financial documents", "Premium commercial rates"],
    # Tier 3
    ["Tier 3", "French ↔ English", "Commercial demand", "Legal, commercial contracts, EU documents", "Widely available but always needed"],
    ["Tier 3", "Spanish ↔ English", "Growing community + business", "Business, legal, NHS", "Steady demand; many translators available"],
    ["Tier 3", "Portuguese (Brazilian) ↔ English", "Growing fast", "NHS, immigration, legal", "Rapidly growing Midlands community"],
    ["Tier 3", "Lithuanian ↔ English", "Baltic community", "Employment, housing, courts", "Consistent Midlands demand"],
]

lang_table = doc.add_table(rows=len(rows_data), cols=5)
lang_table.style = 'Table Grid'
lang_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for r_idx, row_data in enumerate(rows_data):
    row = lang_table.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = col_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(8.5)

        if r_idx == 0:  # Header row
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_shading(cell, '0F2340')
        else:
            if r_idx % 2 == 0:
                set_cell_shading(cell, ROW_GREY_HEX)
            else:
                set_cell_shading(cell, 'FFFFFF')

# ───────────────────────────────────────────────
# SECTION 3 — LEGAL FRAMEWORK
# ───────────────────────────────────────────────

add_heading(doc, "3. Legal Considerations Before Engaging Freelancers", gold_border=True)

add_heading2(doc, "3.1 Employment Status — The Three Categories")
add_body(doc, "UK law recognises three work statuses: Employee, Worker, and Self-Employed Independent Contractor. Metaphrase's translators will be engaged as self-employed independent contractors — which is correct and legal, provided the relationship genuinely reflects this. A genuine self-employed contractor: works for multiple clients, is not obliged to accept work, uses their own equipment, bears their own financial risk, and is paid per project with no holiday pay or pension. The Subcontractor Agreement signed in Phase 1 correctly defines this relationship. Never treat a contractor as self-employed if in practice they work exclusively for Metaphrase, follow fixed hours, or are integrated into the business as if an employee — this is 'bogus self-employment' and exposes Metaphrase to tribunal claims for holiday pay, national minimum wage, and auto-enrolment pension contributions.")

add_heading2(doc, "3.2 IR35 / Off-Payroll Working Rules")
add_body(doc, "IR35 concerns contractors who provide services through their own limited company (a Personal Service Company or PSC). For individual freelancers working under their own name, IR35 does not apply. However, if a subcontractor operates through their own limited company, Metaphrase may need to issue a Status Determination Statement (SDS) assessing whether the engagement would constitute employment. For small businesses like Metaphrase, use HMRC's free CEST tool (Check Employment Status for Tax) at gov.uk/guidance/check-employment-status-for-tax. In practice, most freelance translators work as sole traders, not through limited companies, so IR35 will rarely be relevant at this stage.")

add_heading2(doc, "3.3 Right to Work")
add_body(doc, "Metaphrase is not legally required to conduct formal Right to Work checks for self-employed subcontractors in the same way as for employees. However, it is strongly recommended best practice to: (a) request a copy of each subcontractor's passport or national identity document; (b) note that you have verified their right to undertake paid work in the UK; and (c) retain a record of this verification. Engaging someone who does not have the right to work — even as a self-employed contractor — could expose Metaphrase to reputational and legal risk. Keep copies on file.")

add_heading2(doc, "3.4 GDPR — Processing Applicant Data")
add_body(doc, "Collecting CVs, qualifications, and personal details from applicants means processing personal data under UK GDPR. Obligations: (a) display a privacy notice on the Work With Us page before the form, explaining what data is collected, why, retention period, and who can access it; (b) retain unsuccessful applications for no more than 12 months then delete; (c) restrict Google Sheet access to authorised users only; (d) if storing documents in Google Drive, use folder-level permissions. Metaphrase's existing ICO registration covers this provided the processing is disclosed in the privacy notice.")

add_heading2(doc, "3.5 What Qualifications to Require (and How to Verify)")
add_body(doc, "Require and verify the following credentials:")
add_bullet(doc, "CIOL membership — most prestigious; verify at ciol.org.uk/member-check using member number")
add_bullet(doc, "ITI membership — verify at iti.org.uk/find-a-member")
add_bullet(doc, "NRPSI registration — mandatory for police and court interpreting; verify at nrpsi.org.uk")
add_bullet(doc, "Relevant degree — request copy of certificate; do not take applicants at their word")
add_bullet(doc, "DBS check — Enhanced DBS required for any work involving NHS patients or criminal justice settings; ask applicants if they hold one")
add_bullet(doc, "Professional Indemnity insurance — request copy of certificate as part of onboarding")

add_heading2(doc, "3.6 Paying Subcontractors — Tax")
add_body(doc, "Subcontractors invoice Metaphrase and are responsible for their own Self Assessment income tax and National Insurance. Metaphrase does not deduct PAYE tax at source. If a subcontractor is VAT registered, they will add VAT to their invoice — Metaphrase can reclaim this as input VAT (once VAT registered itself). Keep all subcontractor invoices for 6 years.")

add_heading2(doc, "3.7 Insurance")
add_body(doc, "Every subcontractor should hold their own Professional Indemnity insurance (minimum £500,000 recommended). Metaphrase's own £1 million PI policy protects Metaphrase's liability to clients — it does not extend to cover errors made by subcontractors that Metaphrase fails to catch. Requiring subcontractors to hold PI insurance creates a second layer of protection and demonstrates professionalism to corporate clients.")

# ───────────────────────────────────────────────
# SECTION 4 — THE RECRUITMENT FORM
# ───────────────────────────────────────────────

add_heading(doc, "4. The Translator & Interpreter Application Form", gold_border=True)

add_heading2(doc, "4.1 Platform")
add_body(doc, "Use Google Forms — free, unlimited responses, automatically populates a linked Google Sheet, and supports file uploads (CV, qualification documents). The form is embedded on the website at metaphrase.uk/work-with-us/.")

add_heading2(doc, "4.2 Form Questions")
add_body(doc, "Create the form at forms.google.com. Organise it into the following sections:")

form_sections = [
    ("SECTION A — Personal Details", "Full legal name * | Email address * | Phone number * | City/Town * | Country of residence * | UK-based? (Yes / No)"),
    ("SECTION B — Languages", "Primary language pair (e.g. Urdu → English) * | Additional language pairs | Native speaker of source language? (Yes / No) | Target language level (Native / Near-native / Professional working)"),
    ("SECTION C — Specialisations (tick all that apply)", "Legal & Courts | Immigration & UKVI | Medical & NHS | Financial & Business | Social Services | Education & Schools | General/Administrative | Conference & Events | Other (specify)"),
    ("SECTION D — Qualifications & Memberships", "Highest relevant qualification and institution | Professional memberships: CIOL / ITI / NRPSI / Other | CIOL or ITI membership number (for verification) | DBS checked? (Enhanced / Basic / Not currently)"),
    ("SECTION E — Experience", "Years of experience * | Worked for a translation agency before? (Yes / No) | Types of work: Document translation / Telephone interpreting / Face-to-face interpreting / Court interpreting / Conference / Subtitling / Transcription"),
    ("SECTION F — Availability", "Typical weekly capacity (hours) | Can accept urgent/same-day requests? (Yes / Sometimes / No) | Secure remote working setup for video/telephone interpreting? (Yes / No)"),
    ("SECTION G — Rates & Tax", "Expected rate per word for translation (£) | Expected rate per hour for interpreting (£) | VAT registered? (Yes / No) | Invoice through limited company? (Yes / No)"),
    ("SECTION H — Tools", "CAT software: SDL Trados / memoQ / Wordfast / OmegaT / None / Other | Microsoft Office level (Basic / Intermediate / Advanced)"),
    ("SECTION I — Compliance & Agreements", "Professional Indemnity insurance? (Yes / No / Willing to obtain) | Legal right to work in the UK? (Yes / No) * | I agree to maintain strict confidentiality for all work * (checkbox) | I agree to Metaphrase's processing of my data for recruitment * (checkbox)"),
    ("SECTION J — Supporting Documents", "Upload CV or portfolio (file) | Upload evidence of qualifications/membership (file)"),
    ("SECTION K — How Did You Hear About Us?", "LinkedIn / ProZ / CIOL / Google / Referred by someone / Other"),
]

for section_title, section_content in form_sections:
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(6)
    p_sec.paragraph_format.space_after = Pt(2)
    r_sec = p_sec.add_run(section_title)
    r_sec.bold = True
    r_sec.font.size = Pt(9.5)
    r_sec.font.color.rgb = BLUE

    p_cont = doc.add_paragraph()
    p_cont.paragraph_format.space_before = Pt(0)
    p_cont.paragraph_format.space_after = Pt(3)
    p_cont.paragraph_format.left_indent = Cm(0.5)
    r_cont = p_cont.add_run(section_content)
    r_cont.font.size = Pt(9)

add_heading2(doc, "4.3 Google Sheet Setup")
add_body(doc, "Google Forms auto-creates a linked spreadsheet. Add a \"Status\" column manually with values: New / Under Review / Test Sent / Approved / Rejected. Use conditional formatting: red fill for rows where Right to Work = No; amber for PI Insurance = No; green for rows marked Approved.")

# ───────────────────────────────────────────────
# SECTION 5 — WHERE TO SHARE THE FORM
# ───────────────────────────────────────────────

add_heading(doc, "5. Where to Share the Recruitment Form", gold_border=True)

platforms = [
    ("ProZ.com", "World's largest translator marketplace. Post a \"Translator Sought\" listing in each target language pair specifying language pair, specialisation, rate range, and link to the application form. Free to post in forums; £0–£50 depending on visibility options."),
    ("TranslatorsCafé.com", "Second major translator marketplace. Post in the Job Offers section. Free."),
    ("ITI (Institute of Translation and Interpreting)", "Post in the ITI job board and members' forum. Direct-message ITI members whose public profiles match target language pairs. Membership required (Kashif can post as a CIOL member with equivalent standing)."),
    ("CIOL Network", "As a CIOL member, email fellow CIOL members directly in target language pairs — a personal introduction carries far more weight than a job board posting. Access via the CIOL member directory."),
    ("NRPSI Register", "For court and public service interpreters specifically. Browse the register at nrpsi.org.uk and contact interpreters in target language pairs directly."),
    ("LinkedIn", "Post on the Metaphrase Ltd company page. Search \"freelance Urdu translator UK\", \"Arabic interpreter Birmingham\" etc. and connect personally. Join and post in groups: \"Translators and Interpreters UK\", \"CIOL Members Network\", \"ITI Community\". Free."),
    ("Facebook Groups", "Post in: \"Translators UK\", \"Arabic Translators UK\", \"Polish Translators UK & Europe\", \"Urdu Translators & Interpreters\", \"Somali Translators\", \"Romanian Translators UK\". Most groups are free to post in after approval."),
    ("University Language Departments", "Email course leaders at: University of Birmingham (Modern Languages), Aston University, University of Westminster (MA Translation), University of Bath, Heriot-Watt University (leading translation MA), Leeds Beckett. Offer to be listed as a graduate employer/partner. Free; builds long-term talent pipeline."),
    ("Indeed UK / Reed.co.uk", "Post as a self-employed/freelance contractor role (be explicit it is not employment). Builds Metaphrase's employer profile simultaneously. £0–£100 depending on promoted options."),
    ("The Metaphrase Website", "metaphrase.uk/work-with-us/ with the embedded Google Form. Include in footer navigation. Every person who finds the website organically via SEO can also find the recruitment page."),
]

for platform_name, platform_desc in platforms:
    p_plat = doc.add_paragraph()
    p_plat.paragraph_format.space_before = Pt(6)
    p_plat.paragraph_format.space_after = Pt(2)
    r_bold = p_plat.add_run(platform_name + " — ")
    r_bold.bold = True
    r_bold.font.size = Pt(10)
    r_bold.font.color.rgb = NAVY
    r_desc = p_plat.add_run(platform_desc)
    r_desc.font.size = Pt(10)

# ───────────────────────────────────────────────
# SECTION 6 — VETTING AND ONBOARDING
# ───────────────────────────────────────────────

add_heading(doc, "6. Vetting & Onboarding", gold_border=True)

steps = [
    ("Step 1 — Application Review (Days 1–3 after receipt)", "Screen Google Sheet. Filter: UK-based, relevant language pair, has qualifications, right to work confirmed. Mark status as \"Under Review\"."),
    ("Step 2 — Credential Verification (Days 3–5)", "Verify CIOL membership at ciol.org.uk/member-check. Verify ITI at iti.org.uk/find-a-member. Check NRPSI for interpreters at nrpsi.org.uk. Request degree certificate if no CIOL/ITI membership. Mark verified applicants \"Test Sent\"."),
    ("Step 3 — Test Piece (Days 5–10)", "Send a 200–300 word test passage in the applicant's claimed specialism. PAY for the test at their stated rate — requesting unpaid tests is considered unethical in the translation profession and will deter the best candidates. Assess: accuracy, register, formatting, turnaround time, and communication."),
    ("Step 4 — Brief Call (Optional, 15 minutes)", "For interpreters especially. Assess professionalism, communication, and understanding of ethics (impartiality, confidentiality, no personal opinions during assignments)."),
    ("Step 5 — Onboarding", "Send the Subcontractor Agreement for e-signature (DocuSign free tier or simply PDF via email). Collect: signed agreement, copy of passport/ID, proof of PI insurance, bank details. Add to master database. Mark status \"Approved\"."),
    ("Step 6 — First Project", "Assign a small, low-risk project first. Review output carefully. Build the relationship gradually before assigning complex or high-value work."),
]

for step_title, step_desc in steps:
    p_step = doc.add_paragraph()
    p_step.paragraph_format.space_before = Pt(8)
    p_step.paragraph_format.space_after = Pt(2)
    r_st = p_step.add_run(step_title)
    r_st.bold = True
    r_st.font.size = Pt(10.5)
    r_st.font.color.rgb = NAVY
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after = Pt(4)
    p_desc.paragraph_format.left_indent = Cm(0.5)
    r_d = p_desc.add_run(step_desc)
    r_d.font.size = Pt(10)

# ───────────────────────────────────────────────
# SECTION 7 — RATE STRUCTURE
# ───────────────────────────────────────────────

add_heading(doc, "7. Suggested Rate Structure", gold_border=True)

add_body(doc, "Rates below reflect current UK market conditions (2026). Always negotiate based on language rarity, specialism, and translator seniority. Aim for 40–50% gross margin on all subcontracted work.")

# Rate table
rate_headers = ["Service", "Pay to Subcontractor", "Charge to Client", "Gross Margin"]
rate_rows = [
    ["Translation — common pairs (Polish, Romanian, French, Spanish)", "£0.07–0.10/word", "£0.12–0.18/word", "~40–50%"],
    ["Translation — mid-demand (Urdu, Hindi, Punjabi, Arabic, Bengali)", "£0.08–0.12/word", "£0.15–0.22/word", "~40–50%"],
    ["Translation — rare/specialist (Somali, Pashto, Dari, Gujarati)", "£0.10–0.16/word", "£0.20–0.30/word", "~40–50%"],
    ["Interpreting — general community", "£20–30/hr", "£45–70/hr", "~50%"],
    ["Interpreting — legal/court", "£30–45/hr", "£70–120/hr", "~50%"],
    ["Interpreting — medical/NHS", "£25–35/hr", "£55–85/hr", "~50%"],
]

rate_col_widths = [Cm(6.5), Cm(3.0), Cm(3.0), Cm(2.5)]
rate_table = doc.add_table(rows=len(rate_rows)+1, cols=4)
rate_table.style = 'Table Grid'
rate_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for c_idx, hdr in enumerate(rate_headers):
    cell = rate_table.rows[0].cells[c_idx]
    cell.width = rate_col_widths[c_idx]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = WHITE
    set_cell_shading(cell, '0F2340')

for r_idx, row_data in enumerate(rate_rows):
    row = rate_table.rows[r_idx + 1]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = rate_col_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9)
        if r_idx % 2 == 0:
            set_cell_shading(cell, ROW_GREY_HEX)
        else:
            set_cell_shading(cell, 'FFFFFF')

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(6)
p_note.paragraph_format.space_after = Pt(4)
p_note.paragraph_format.left_indent = Cm(0.3)
r_note = p_note.add_run("Note: Never undercut the subcontractor below sustainable rates — the best linguists will decline and go elsewhere. Competitive pay attracts quality, and quality retains clients.")
r_note.font.size = Pt(9)
r_note.italic = True
r_note.font.color.rgb = GREY

# ───────────────────────────────────────────────
# SECTION 8 — WORK WITH US PAGE
# ───────────────────────────────────────────────

add_heading(doc, "8. Adding 'Work With Us' to the Website", gold_border=True)

website_steps = [
    ("8.1", "The page is live at metaphrase.uk/work-with-us/ (created as part of this Phase 2 delivery)."),
    ("8.2", "Once your Google Form is created, copy the embed URL from Google Forms (click Send → embed tab → copy the iframe src URL) and replace the placeholder in the page HTML. The file is at: metaphrase-main/work-with-us/index.html."),
    ("8.3", "Add \"Work With Us\" to the footer of every page under the Company column."),
    ("8.4", "The page has been added to sitemap.xml."),
    ("8.5", "Share the direct URL (metaphrase.uk/work-with-us/) in all your ProZ, LinkedIn, and Facebook recruitment posts."),
]

for step_num, step_text in website_steps:
    p_ws = doc.add_paragraph()
    p_ws.paragraph_format.space_before = Pt(4)
    p_ws.paragraph_format.space_after = Pt(4)
    r_num = p_ws.add_run(f"{step_num} — ")
    r_num.bold = True
    r_num.font.size = Pt(10)
    r_num.font.color.rgb = NAVY
    r_txt = p_ws.add_run(step_text)
    r_txt.font.size = Pt(10)

# ───────────────────────────────────────────────
# SECTION 9 — PRIVACY NOTICE
# ───────────────────────────────────────────────

add_heading(doc, "9. Privacy Notice (Display on Work With Us Page)", gold_border=True)

add_body(doc, "This notice must appear on the page before the application form. The following text is ready to use:")

# Shaded box paragraph for privacy notice
p_priv = doc.add_paragraph()
p_priv.paragraph_format.space_before = Pt(6)
p_priv.paragraph_format.space_after = Pt(6)
p_priv.paragraph_format.left_indent = Cm(0.5)
p_priv.paragraph_format.right_indent = Cm(0.5)
set_para_shading(p_priv, 'EFF3FF')
r_priv = p_priv.add_run(
    "Privacy Notice — Translator & Interpreter Applications. Metaphrase Ltd (Company No. 16371031) is the data controller for information you submit in this application. "
    "We collect your personal information solely to assess your suitability as a freelance subcontractor. Your data will be kept securely, accessed only by authorised Metaphrase staff, and retained for 12 months if your application is unsuccessful, after which it will be permanently deleted. "
    "If your application is successful, your data will be retained for the duration of our working relationship and for 6 years thereafter in accordance with UK record-keeping requirements. "
    "You have the right to access, correct, or request deletion of your data at any time by emailing info@metaphrase.uk. Metaphrase Ltd is registered with the Information Commissioner's Office (ICO)."
)
r_priv.font.size = Pt(9.5)
r_priv.italic = True

# ───────────────────────────────────────────────
# SECTION 10 — 90-DAY ACTION PLAN
# ───────────────────────────────────────────────

add_heading(doc, "10. 90-Day Action Plan", gold_border=True)

action_headers = ["Week", "Action", "Notes", "Done?"]
action_rows = [
    ["Week 1", "Create Google Form using questions in Section 4", "Use forms.google.com", "[ ]"],
    ["Week 1", "Ensure Work With Us page is live", "Already built in Phase 2", "[✓]"],
    ["Week 1", "Add Google Form embed URL to work-with-us page", "Replace placeholder src", "[ ]"],
    ["Week 1", "Add \"Work With Us\" to website footer navigation", "Already done", "[✓]"],
    ["Week 2", "Post on ProZ.com for Tier 1 language pairs", "Create 6 separate listings", "[ ]"],
    ["Week 2", "Post on TranslatorsCafé", "Mirror ProZ listings", "[ ]"],
    ["Week 2", "Post on LinkedIn company page and personal profile", "Tag relevant language communities", "[ ]"],
    ["Week 2", "Email 5 CIOL members directly in Tier 1 pairs", "Personal email; mention shared CIOL membership", "[ ]"],
    ["Week 3", "Post in Facebook translator groups", "5–6 relevant groups", "[ ]"],
    ["Week 3", "Email language departments at 3 universities", "Birmingham, Aston, Westminster", "[ ]"],
    ["Week 3", "Post on Indeed / Reed as self-employed contractor role", "Be explicit: not employment", "[ ]"],
    ["Week 4", "Review first batch of applications", "Screen Google Sheet; mark status", "[ ]"],
    ["Week 5–6", "Credential verification for shortlisted applicants", "Use CIOL/ITI/NRPSI online checks", "[ ]"],
    ["Week 5–6", "Send paid test translations to qualified applicants", "200–300 words; pay at their rate", "[ ]"],
    ["Week 6–8", "Sign Subcontractor Agreements", "Use Phase 1 agreement template", "[ ]"],
    ["Week 8", "Target: 2+ vetted translators per Tier 1 language pair", "12 approved subcontractors minimum", "[ ]"],
    ["Week 9–10", "Begin Tier 2 language pair recruitment", "Repeat process for Somali, Bengali, Gujarati etc.", "[ ]"],
    ["Week 12", "Review coverage; identify remaining gaps", "Assess against demand pipeline", "[ ]"],
]

action_col_widths = [Cm(1.8), Cm(6.0), Cm(5.5), Cm(1.7)]
action_table = doc.add_table(rows=len(action_rows)+1, cols=4)
action_table.style = 'Table Grid'
action_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for c_idx, hdr in enumerate(action_headers):
    cell = action_table.rows[0].cells[c_idx]
    cell.width = action_col_widths[c_idx]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = WHITE
    set_cell_shading(cell, '0F2340')

for r_idx, row_data in enumerate(action_rows):
    row = action_table.rows[r_idx + 1]
    for c_idx, text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = action_col_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9)
        if r_idx % 2 == 0:
            set_cell_shading(cell, ROW_GREY_HEX)
        else:
            set_cell_shading(cell, 'FFFFFF')

# Final paragraph
p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(20)
p_end.paragraph_format.space_after = Pt(8)
r_end = p_end.add_run("— End of Document —")
r_end.font.size = Pt(10)
r_end.italic = True
r_end.font.color.rgb = GREY
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ───────────────────────────────────────────────
# SAVE
# ───────────────────────────────────────────────
output_path = "/sessions/dreamy-hopeful-mendel/mnt/metaphrase-main/docs/Metaphrase-Phase2-Plan.docx"
doc.save(output_path)
print(f"SUCCESS: Document saved to {output_path}")
